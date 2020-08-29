import re
import datetime
import docx #ref: https://python-docx.readthedocs.io/en/latest/

def main():
    ''' map store all input keys'''
    dic = {'body1': "", 'body2': "", 'user': "", 'userAddress1': "", 'userAddress2': "",
           'contactNumber': "", 'email': "", 'compName': "", 'compAddress1': "",
           'compAddress2': "", 'position': "", 'positionNO': ""}

    ''' store 3 body paragraphs '''
    body = {'intro': "", 'b1': "", 'b2': ""}

    ''' cover letter model '''
    model = ""
    doc1 = docx.Document("model.docx")
    file = doc1.paragraphs

    ''' store paragraphs and seperate each paragraph by an empty line '''
    for line in (file):
        model += line.text
        model += ("\n")

    ''' ask a few question to construct a cover letter '''
    getInfo(dic)

    ''' get body paragraphs from input file '''
    doc2 = docx.Document("input.docx")
    allPara = doc2.paragraphs
    key = ""
    for line in allPara:
        if key == "":
            key = getParagraph(line.text, dic)
            continue

        if key == 'intro':
            body['intro'] += " " + line.text.strip()
            if line.text == "":
                key = ""
                continue

        if key == str(dic['body1']):
            body['b1'] += " " + line.text.strip()
            if line.text == "":
                key = ""
                continue

        if key == str(dic['body2']):
            body['b2'] += " " + line.text.strip()
            if line.text == "":
                key = ""
                continue

    body['intro'] = body['intro'].strip()
    body['intro'] = re.sub(r"== COMP NAME ==", dic['compName'], body['intro'])
    body['intro'] = re.sub(r"== POSITION ==", dic['position'], body['intro'])

    body['b1'] = body['b1'].strip()
    body['b2'] = body['b2'].strip()

    ''' get current date '''
    date = getDate()

    ''' sub in value '''
    f_model = subIn(dic, model, body, date)

    ''' make a cover letter and export it in aa form of docx '''
    constructDoxc(f_model, dic)

''' sub in all info '''
def subIn(dic, model, body, date):
    f_model = re.sub(r"== NAME ==", dic['user'], model)
    f_model = re.sub(r"== ADDRESS 1 ==", dic['userAddress1'], f_model)
    f_model = re.sub(r"== ADDRESS 2 ==", dic['userAddress2'], f_model)
    f_model = re.sub(r"== DATE ==", date, f_model)
    f_model = re.sub(r"== COMP NAME ==", dic['compName'], f_model)
    f_model = re.sub(r"== COMP ADDRESS 1 ==", dic['compAddress1'], f_model)
    f_model = re.sub(r"== COMP ADDRESS 2 ==", dic['compAddress2'], f_model)
    f_model = re.sub(r"== POSITION ==", dic['position'], f_model)
    f_model = re.sub(r"== POSITION NUMBER ==", dic['positionNO'], f_model)
    f_model = re.sub(r"== INTRO ==", body['intro'], f_model)
    f_model = re.sub(r"== BODY 1 ==", body['b1'], f_model)
    f_model = re.sub(r"== BODY 2 ==", body['b2'], f_model)

    return f_model

''' get current date '''
def getDate():
    date = datetime.datetime.now()
    str = date.strftime("%B") + " " + date.strftime("%d") + ", " + date.strftime("%Y")
    return str

''' extract code from the file '''
def getParagraph(line, dic):
    pattern = re.compile(r"^{{ (.+) }}?")
    matchObj = pattern.search(line)

    if matchObj:
        if matchObj.group(1) == "intro":
            return matchObj.group(1)

        if matchObj.group(1) == str(dic['body1']):
            return matchObj.group(1)

        if matchObj.group(1) == str(dic['body2']):
            return matchObj.group(1)
    return ""

''' get info to make a cover letter '''
def getInfo(dic):
    ''' ask a few questions to construct a cover letter '''
    dic['body1'] = input("Enter code of your first body paragraph: ").strip()
    dic['body2'] = input("Enter code of your second body paragraph: ").strip()
    dic['user'] = input("Enter your name: ").strip()
    dic['userAddress1'] = input("Enter your address line 1: ").strip()
    dic['userAddress2'] = input("Enter your address line 2: ").strip()
    dic['contactNumber'] = input("Enter your contact number: ").strip()
    dic['email'] = input("Enter your contact email: ").strip()
    dic['compName'] = input("Enter the company's name that you are applying for: ").strip()
    dic['compAddress1'] = input("Enter the company's address line 1: ").strip()
    dic['compAddress2'] = input("Enter the company's address line 2: ").strip()
    dic['position'] = input("Enter the position you are applying for: ").strip()
    dic['positionNO'] = input("Enter the job/position number if it is applicable: ").strip()

    ''' format code for paragraph '''
    dic['body1'] = re.sub(r"[\D]+", '', dic['body1'])
    dic['body2'] = re.sub(r"[\D]+", '', dic['body2'])

    ''' format phone number '''
    num = re.sub(r"[\D]+", '', dic['contactNumber'])
    dic['contactNumber'] = "(" + num[:3] + ")-" + num[3:6] + "-" + num[6:]

    #format mailing address using regex

''' construct a docx file and export it '''
def constructDoxc(text, dic):
    ''' get every line store in a list '''
    line = []
    line = re.split(r'[\n\r]', text)

    ''' create mydoc obj '''
    mydoc = docx.Document()

    ''' create style obj (access by document.styles) '''
    style = mydoc.styles['Normal']

    ''' modify font (style.class.method) '''
    font = style.font
    font.name = 'Arial'
    font.size = docx.shared.Pt(11)

    ''' modify paragraph (style.class.method) '''
    paragraphFormat = style.paragraph_format
    paragraphFormat.space_after = docx.shared.Pt(0)

    ''' store print each line to docx file '''
    for counter, li in enumerate(line):
        ''' make heading '''
        if counter == 0:

            ''' modify specific text (called run) '''
            ''' docx.text.run.FONT '''
            head = mydoc.add_paragraph().add_run(li)
            head.bold = True
            head.font.size = docx.shared.Pt(18)
            continue

        #make these 3 for...loop better use regex to extract the condition
        ''' format contact number '''
        if counter == 3:
            formatAilgn(mydoc, li, dic['contactNumber'])
            continue

        ''' format email '''
        if counter == 4:
            formatAilgn(mydoc, li, dic['email'])
            continue

        if counter == 8:
            head = mydoc.add_paragraph().add_run(li)
            head.bold = True
            continue
        mydoc.add_paragraph(li)

    ''' save and export '''
    mydoc.save("output.docx")

''' Big thanks for the reference: https://stackoverflow.com/questions/58656450/how-to-use-tabletop-by-python-docx '''
''' format left and right alian in the same line '''
def formatAilgn(mydoc, currLi, text):
    p = mydoc.add_paragraph(f'{currLi}\t{text}')  # tab will trigger tabstop
    sec = mydoc.sections[0]

    ''' finding end_point for the content '''
    margin_end = docx.shared.Inches(sec.page_width.inches - (sec.left_margin.inches + sec.right_margin.inches))
    tab_stops = p.paragraph_format.tab_stops

    ''' adding new tab stop, to the end point, and making sure that it's `RIGHT` aligned. '''
    tab_stops.add_tab_stop(margin_end, docx.enum.text.WD_TAB_ALIGNMENT.RIGHT)

''' call main() '''
if __name__ == "__main__":
    main()